"""
File Summarizer Tool

This tool reads various file formats (PDF, DOCX, TXT, MD, etc.), extracts content,
and generates AI-powered summaries using AWS Bedrock.

Features:
- Multi-format file reading (PDF, DOCX, TXT, MD, etc.)
- File content extraction and preprocessing
- Summary generation via AWS Bedrock
- Summary length control (short/standard/detailed)
- Output format selection (json/markdown/plain)
- Key information extraction
- Batch file processing
- Error handling and exception management

Dependencies:
- PyPDF2/pdfplumber: PDF processing
- python-docx: Word document processing
- markdown: Markdown processing
- chardet: Encoding detection
- boto3: AWS SDK for Python
"""

import os
import json
import time
import logging
import tempfile
from pathlib import Path
from typing import Dict, List, Union, Optional, Any, Tuple, Literal
from concurrent.futures import ThreadPoolExecutor, as_completed
import re
import traceback

# Import necessary libraries for file processing
try:
    import PyPDF2
    import pdfplumber
    import docx
    import markdown
    import chardet
    import boto3
except ImportError as e:
    missing_lib = str(e).split("'")[1]
    raise ImportError(f"Missing required library: {missing_lib}. Please install it using pip.")

from strands import tool

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger('file_summarizer_tool')

# Define summary length parameters
SUMMARY_LENGTH_CONFIG = {
    'short': {
        'max_tokens': 100,
        'description': 'Brief overview with core points only (50-100 words)'
    },
    'standard': {
        'max_tokens': 250,
        'description': 'Comprehensive summary with main points and supporting details (150-300 words)'
    },
    'detailed': {
        'max_tokens': 500,
        'description': 'In-depth summary covering all significant aspects (300-500 words)'
    }
}

# File type handlers
class FileHandler:
    """Base class for file handlers"""
    
    @staticmethod
    def detect_encoding(file_path: str) -> str:
        """Detect file encoding"""
        with open(file_path, 'rb') as file:
            raw_data = file.read()
            result = chardet.detect(raw_data)
            return result['encoding'] or 'utf-8'
    
    @staticmethod
    def chunk_text(text: str, chunk_size: int = 4000) -> List[str]:
        """Split text into manageable chunks"""
        if len(text) <= chunk_size:
            return [text]
        
        # Try to split at paragraph breaks
        paragraphs = text.split('\n\n')
        chunks = []
        current_chunk = ""
        
        for paragraph in paragraphs:
            if len(current_chunk) + len(paragraph) + 2 <= chunk_size:
                if current_chunk:
                    current_chunk += "\n\n"
                current_chunk += paragraph
            else:
                if current_chunk:
                    chunks.append(current_chunk)
                
                # If paragraph is longer than chunk_size, split it
                if len(paragraph) > chunk_size:
                    words = paragraph.split(' ')
                    current_chunk = ""
                    for word in words:
                        if len(current_chunk) + len(word) + 1 <= chunk_size:
                            if current_chunk:
                                current_chunk += " "
                            current_chunk += word
                        else:
                            chunks.append(current_chunk)
                            current_chunk = word
                else:
                    current_chunk = paragraph
        
        if current_chunk:
            chunks.append(current_chunk)
            
        return chunks


class PDFHandler(FileHandler):
    """Handler for PDF files"""
    
    @staticmethod
    def extract_content(file_path: str) -> str:
        """Extract text content from PDF file"""
        text = ""
        metadata = {}
        
        # First try with PyPDF2
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                metadata = {
                    'pages': len(reader.pages),
                    'title': reader.metadata.get('/Title', ''),
                    'author': reader.metadata.get('/Author', ''),
                    'subject': reader.metadata.get('/Subject', ''),
                    'creator': reader.metadata.get('/Creator', '')
                }
                
                for page_num in range(len(reader.pages)):
                    page = reader.pages[page_num]
                    text += page.extract_text() + "\n\n"
        except Exception as e:
            logger.warning(f"PyPDF2 extraction failed: {str(e)}. Trying pdfplumber...")
            
            # If PyPDF2 fails, try with pdfplumber
            try:
                with pdfplumber.open(file_path) as pdf:
                    metadata = {
                        'pages': len(pdf.pages)
                    }
                    
                    for page in pdf.pages:
                        page_text = page.extract_text() or ""
                        text += page_text + "\n\n"
            except Exception as e2:
                raise ValueError(f"Failed to extract text from PDF: {str(e2)}")
        
        return text.strip(), metadata


class DocxHandler(FileHandler):
    """Handler for DOCX files"""
    
    @staticmethod
    def extract_content(file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text content from DOCX file"""
        try:
            doc = docx.Document(file_path)
            
            # Extract metadata
            metadata = {
                'title': doc.core_properties.title or '',
                'author': doc.core_properties.author or '',
                'subject': doc.core_properties.subject or '',
                'created': str(doc.core_properties.created) if doc.core_properties.created else '',
                'modified': str(doc.core_properties.modified) if doc.core_properties.modified else '',
                'paragraphs': len(doc.paragraphs)
            }
            
            # Extract full text
            text = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            
            return text, metadata
            
        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")


class TextHandler(FileHandler):
    """Handler for plain text files"""
    
    @staticmethod
    def extract_content(file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract content from text file"""
        try:
            encoding = FileHandler.detect_encoding(file_path)
            with open(file_path, 'r', encoding=encoding) as file:
                text = file.read()
            
            lines = text.split('\n')
            metadata = {
                'lines': len(lines),
                'encoding': encoding
            }
            
            return text, metadata
            
        except Exception as e:
            raise ValueError(f"Failed to extract text from text file: {str(e)}")


class MarkdownHandler(FileHandler):
    """Handler for Markdown files"""
    
    @staticmethod
    def extract_content(file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract content from Markdown file"""
        try:
            encoding = FileHandler.detect_encoding(file_path)
            with open(file_path, 'r', encoding=encoding) as file:
                md_text = file.read()
            
            # Extract metadata if available (YAML frontmatter)
            metadata = {}
            frontmatter_match = re.match(r'^---\s*\n(.*?)\n---\s*\n', md_text, re.DOTALL)
            if frontmatter_match:
                try:
                    import yaml
                    frontmatter = frontmatter_match.group(1)
                    metadata = yaml.safe_load(frontmatter)
                    # Remove frontmatter from text
                    md_text = re.sub(r'^---\s*\n.*?\n---\s*\n', '', md_text, flags=re.DOTALL)
                except (ImportError, yaml.YAMLError):
                    pass
            
            # Extract headers for structure
            headers = re.findall(r'^(#+)\s+(.+)$', md_text, re.MULTILINE)
            if headers:
                metadata['headers'] = [{'level': len(h[0]), 'text': h[1]} for h in headers]
            
            # Add basic metadata
            metadata['lines'] = len(md_text.split('\n'))
            metadata['encoding'] = encoding
            
            # Convert markdown to plain text for better summarization
            plain_text = md_text
            # Remove code blocks
            plain_text = re.sub(r'```.*?```', '', plain_text, flags=re.DOTALL)
            # Remove inline code
            plain_text = re.sub(r'`.*?`', '', plain_text)
            # Remove links but keep text
            plain_text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', plain_text)
            # Remove images
            plain_text = re.sub(r'!\[.*?\]\(.*?\)', '', plain_text)
            
            return plain_text, metadata
            
        except Exception as e:
            raise ValueError(f"Failed to extract text from Markdown file: {str(e)}")


# File handler factory
class FileHandlerFactory:
    """Factory for creating appropriate file handlers"""
    
    @staticmethod
    def get_handler(file_path: str) -> Tuple[Any, str]:
        """Get appropriate handler for the file type"""
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.pdf':
            return PDFHandler, 'pdf'
        elif file_extension == '.docx':
            return DocxHandler, 'docx'
        elif file_extension == '.md':
            return MarkdownHandler, 'markdown'
        elif file_extension == '.txt':
            return TextHandler, 'text'
        else:
            # Try to handle as text file
            try:
                handler = TextHandler
                handler.extract_content(file_path)
                return TextHandler, 'text'
            except:
                raise ValueError(f"Unsupported file type: {file_extension}")


# AWS Bedrock integration for summary generation
class BedrockSummarizer:
    """Class for generating summaries using AWS Bedrock"""
    
    def __init__(self, model_id: str = "anthropic.claude-3-sonnet-20240229-v1:0"):
        """Initialize the summarizer with specified model"""
        self.bedrock_runtime = boto3.client('bedrock-runtime')
        self.model_id = model_id
    
    def generate_summary(
        self, 
        text: str, 
        summary_length: str = 'standard',
        extract_keywords: bool = False
    ) -> Dict[str, Any]:
        """Generate summary using AWS Bedrock"""
        # Prepare the prompt based on summary length
        length_config = SUMMARY_LENGTH_CONFIG.get(summary_length, SUMMARY_LENGTH_CONFIG['standard'])
        
        # Prepare system prompt
        system_prompt = f"""You are an expert document summarizer. Your task is to create a {summary_length} summary of the provided document.
Follow these guidelines:
- Create a {length_config['description']}
- Identify and include the main topics and key points
- Maintain the original meaning and intent
- Use clear, concise language
- Organize the summary logically
- Include a title that captures the document's essence
"""
        
        if extract_keywords:
            system_prompt += "- Extract 5-10 keywords that best represent the document content\n"
        
        # Prepare user prompt
        user_prompt = f"""Please summarize the following document:

{text}

Please format your response as a structured summary with:
1. A concise title
2. The main summary
3. 3-5 key points as bullet points
"""
        
        if extract_keywords:
            user_prompt += "4. 5-10 relevant keywords"
        
        # Prepare request body
        request_body = {
            "anthropic_version": "bedrock-2023-05-31",
            "max_tokens": length_config['max_tokens'],
            "temperature": 0.1,
            "system": system_prompt,
            "messages": [
                {
                    "role": "user",
                    "content": user_prompt
                }
            ]
        }
        
        # Make API call to Bedrock
        try:
            response = self.bedrock_runtime.invoke_model(
                modelId=self.model_id,
                body=json.dumps(request_body)
            )
            
            response_body = json.loads(response.get('body').read().decode('utf-8'))
            summary_text = response_body['content'][0]['text']
            
            # Extract title, summary, key points, and keywords
            title_match = re.search(r'#\s+(.*?)(?:\n|$)', summary_text)
            title = title_match.group(1) if title_match else "Summary"
            
            # Extract main summary (text between title and key points)
            main_summary = re.search(r'(?:#\s+.*?\n)(.*?)(?:\n\s*(?:Key|Main) Points|\Z)', summary_text, re.DOTALL)
            summary = main_summary.group(1).strip() if main_summary else summary_text
            
            # Extract key points
            key_points_section = re.search(r'(?:Key|Main) Points[:\n]+(.*?)(?:\n\s*(?:Keywords|$)|\Z)', summary_text, re.DOTALL)
            key_points = []
            if key_points_section:
                key_points_text = key_points_section.group(1)
                key_points = [point.strip().lstrip('*-•⁃◦‣⁌⁍').strip() for point in key_points_text.split('\n') if point.strip()]
            
            # Extract keywords if requested
            keywords = []
            if extract_keywords:
                keywords_section = re.search(r'Keywords[:\n]+(.*?)(?:\n\s*$|\Z)', summary_text, re.DOTALL)
                if keywords_section:
                    keywords_text = keywords_section.group(1)
                    keywords = [kw.strip().strip(',').strip() for kw in re.split(r'[,\n]', keywords_text) if kw.strip()]
            
            return {
                "title": title,
                "summary": summary,
                "key_points": key_points,
                "keywords": keywords if extract_keywords else []
            }
            
        except Exception as e:
            logger.error(f"Error generating summary: {str(e)}")
            raise RuntimeError(f"Failed to generate summary: {str(e)}")
    
    def process_large_text(
        self, 
        text: str, 
        summary_length: str = 'standard',
        extract_keywords: bool = False,
        chunk_size: int = 8000
    ) -> Dict[str, Any]:
        """Process large text by chunking and summarizing"""
        if len(text) <= chunk_size:
            return self.generate_summary(text, summary_length, extract_keywords)
        
        # Split text into chunks
        handler = FileHandler()
        chunks = handler.chunk_text(text, chunk_size)
        
        # Generate summary for each chunk
        chunk_summaries = []
        for i, chunk in enumerate(chunks):
            logger.info(f"Processing chunk {i+1}/{len(chunks)}")
            chunk_summary = self.generate_summary(chunk, summary_length, False)
            chunk_summaries.append(chunk_summary["summary"])
        
        # Combine chunk summaries and generate final summary
        combined_text = "\n\n".join(chunk_summaries)
        final_summary = self.generate_summary(combined_text, summary_length, extract_keywords)
        
        return final_summary


# Output formatters
class OutputFormatter:
    """Format summary output in different formats"""
    
    @staticmethod
    def format_json(summary_data: Dict[str, Any], metadata: Dict[str, Any]) -> str:
        """Format summary as JSON"""
        output = {
            "title": summary_data.get("title", "Summary"),
            "summary": summary_data.get("summary", ""),
            "key_points": summary_data.get("key_points", []),
            "metadata": metadata
        }
        
        if "keywords" in summary_data and summary_data["keywords"]:
            output["keywords"] = summary_data["keywords"]
            
        return json.dumps(output, indent=2)
    
    @staticmethod
    def format_markdown(summary_data: Dict[str, Any], metadata: Dict[str, Any]) -> str:
        """Format summary as Markdown"""
        md_output = f"# {summary_data.get('title', 'Summary')}\n\n"
        md_output += f"{summary_data.get('summary', '')}\n\n"
        
        if summary_data.get("key_points"):
            md_output += "## Key Points\n\n"
            for point in summary_data["key_points"]:
                md_output += f"* {point}\n"
            md_output += "\n"
        
        if summary_data.get("keywords"):
            md_output += "## Keywords\n\n"
            md_output += ", ".join(summary_data["keywords"]) + "\n\n"
        
        md_output += "## Document Metadata\n\n"
        md_output += "```\n"
        md_output += json.dumps(metadata, indent=2)
        md_output += "\n```\n"
        
        return md_output
    
    @staticmethod
    def format_plain(summary_data: Dict[str, Any], metadata: Dict[str, Any]) -> str:
        """Format summary as plain text"""
        plain_output = f"{summary_data.get('title', 'Summary')}\n"
        plain_output += f"{'-' * len(summary_data.get('title', 'Summary'))}\n\n"
        plain_output += f"{summary_data.get('summary', '')}\n\n"
        
        if summary_data.get("key_points"):
            plain_output += "Key Points:\n"
            for i, point in enumerate(summary_data["key_points"], 1):
                plain_output += f"{i}. {point}\n"
            plain_output += "\n"
        
        if summary_data.get("keywords"):
            plain_output += "Keywords: " + ", ".join(summary_data["keywords"]) + "\n\n"
        
        plain_output += "Document Metadata:\n"
        for key, value in metadata.items():
            if isinstance(value, dict):
                plain_output += f"{key}:\n"
                for sub_key, sub_value in value.items():
                    plain_output += f"  {sub_key}: {sub_value}\n"
            else:
                plain_output += f"{key}: {value}\n"
        
        return plain_output


# Main file summarizer tool
@tool
def file_summarizer_tool(
    file_path: Union[str, List[str]],
    summary_length: Literal["short", "standard", "detailed"] = "standard",
    output_format: Literal["json", "markdown", "plain"] = "json",
    extract_keywords: bool = False,
    batch_process: bool = False
) -> str:
    """
    Reads files of various formats and generates AI-powered summaries.
    
    Args:
        file_path: Path to file or list of file paths to summarize
        summary_length: Length of summary to generate ('short', 'standard', or 'detailed')
        output_format: Format for the output ('json', 'markdown', or 'plain')
        extract_keywords: Whether to extract keywords from the content
        batch_process: Whether to process multiple files in batch mode
        
    Returns:
        Structured summary of the file content(s) in the specified format
        
    Raises:
        ValueError: If file format is unsupported or file cannot be read
        RuntimeError: If summary generation fails
    """
    start_time = time.time()
    
    # Initialize summarizer
    summarizer = BedrockSummarizer()
    
    # Handle single file or batch processing
    if isinstance(file_path, str):
        file_paths = [file_path]
        is_batch = False
    else:
        file_paths = file_path
        is_batch = True or batch_process
    
    # Validate files exist
    for path in file_paths:
        if not os.path.exists(path):
            raise ValueError(f"File not found: {path}")
    
    # Process files
    results = []
    errors = []
    
    def process_file(file_path: str) -> Dict[str, Any]:
        """Process a single file"""
        try:
            # Get appropriate handler
            handler_class, file_type = FileHandlerFactory.get_handler(file_path)
            
            # Extract content
            logger.info(f"Extracting content from {file_path}")
            content, file_metadata = handler_class.extract_content(file_path)
            
            # Add file info to metadata
            file_info = {
                "file_name": os.path.basename(file_path),
                "file_type": file_type,
                "file_size": os.path.getsize(file_path),
                "last_modified": time.ctime(os.path.getmtime(file_path))
            }
            
            metadata = {**file_info, **file_metadata}
            
            # Generate summary
            logger.info(f"Generating summary for {file_path}")
            summary_data = summarizer.process_large_text(
                content, 
                summary_length=summary_length,
                extract_keywords=extract_keywords
            )
            
            # Add processing metadata
            processing_info = {
                "processing_time": round(time.time() - start_time, 2),
                "summary_length": summary_length,
                "word_count": {
                    "original": len(content.split()),
                    "summary": len(summary_data["summary"].split())
                },
                "compression_ratio": round(len(summary_data["summary"].split()) / max(1, len(content.split())), 2)
            }
            
            metadata["processing"] = processing_info
            
            # Format output
            if output_format == "json":
                formatted_output = OutputFormatter.format_json(summary_data, metadata)
            elif output_format == "markdown":
                formatted_output = OutputFormatter.format_markdown(summary_data, metadata)
            else:  # plain
                formatted_output = OutputFormatter.format_plain(summary_data, metadata)
            
            return {
                "file_path": file_path,
                "success": True,
                "output": formatted_output,
                "summary_data": summary_data,
                "metadata": metadata
            }
            
        except Exception as e:
            logger.error(f"Error processing {file_path}: {str(e)}")
            logger.debug(traceback.format_exc())
            return {
                "file_path": file_path,
                "success": False,
                "error": str(e),
                "error_type": type(e).__name__
            }
    
    # Process files (sequentially or in parallel)
    if is_batch and len(file_paths) > 1:
        # Use ThreadPoolExecutor for parallel processing
        with ThreadPoolExecutor(max_workers=min(5, len(file_paths))) as executor:
            future_to_file = {executor.submit(process_file, path): path for path in file_paths}
            for future in as_completed(future_to_file):
                result = future.result()
                if result["success"]:
                    results.append(result)
                else:
                    errors.append(result)
    else:
        # Process sequentially
        for path in file_paths:
            result = process_file(path)
            if result["success"]:
                results.append(result)
            else:
                errors.append(result)
    
    # Prepare final output
    total_time = round(time.time() - start_time, 2)
    
    if is_batch:
        # Return batch results
        batch_result = {
            "batch_summary": {
                "total_files": len(file_paths),
                "successful": len(results),
                "failed": len(errors),
                "total_processing_time": total_time
            },
            "results": [],
            "errors": []
        }
        
        # Add successful results
        for result in results:
            if output_format == "json":
                # For JSON, parse the output back to object
                try:
                    result_output = json.loads(result["output"])
                except:
                    result_output = result["output"]
            else:
                result_output = result["output"]
                
            batch_result["results"].append({
                "file_path": result["file_path"],
                "output": result_output
            })
        
        # Add errors
        for error in errors:
            batch_result["errors"].append({
                "file_path": error["file_path"],
                "error": error["error"],
                "error_type": error["error_type"],
                "recommendation": get_error_recommendation(error["error_type"], error["error"])
            })
        
        # Format final output based on output_format
        if output_format == "json":
            return json.dumps(batch_result, indent=2)
        elif output_format == "markdown":
            md_output = "# Batch Summary Results\n\n"
            md_output += f"**Total Files:** {batch_result['batch_summary']['total_files']}  \n"
            md_output += f"**Successful:** {batch_result['batch_summary']['successful']}  \n"
            md_output += f"**Failed:** {batch_result['batch_summary']['failed']}  \n"
            md_output += f"**Total Processing Time:** {batch_result['batch_summary']['total_processing_time']} seconds\n\n"
            
            if results:
                md_output += "## Successful Results\n\n"
                for i, result in enumerate(batch_result["results"], 1):
                    md_output += f"### {i}. {os.path.basename(result['file_path'])}\n\n"
                    if isinstance(result["output"], str):
                        md_output += result["output"] + "\n\n"
                    else:
                        md_output += f"**Title:** {result['output'].get('title', 'Summary')}\n\n"
                        md_output += f"{result['output'].get('summary', '')}\n\n"
                        
                        if "key_points" in result["output"]:
                            md_output += "**Key Points:**\n\n"
                            for point in result["output"]["key_points"]:
                                md_output += f"* {point}\n"
                            md_output += "\n"
            
            if errors:
                md_output += "## Errors\n\n"
                for i, error in enumerate(batch_result["errors"], 1):
                    md_output += f"### {i}. {os.path.basename(error['file_path'])}\n\n"
                    md_output += f"**Error:** {error['error']}\n\n"
                    md_output += f"**Recommendation:** {error['recommendation']}\n\n"
            
            return md_output
        else:  # plain
            plain_output = "BATCH SUMMARY RESULTS\n"
            plain_output += "=====================\n\n"
            plain_output += f"Total Files: {batch_result['batch_summary']['total_files']}\n"
            plain_output += f"Successful: {batch_result['batch_summary']['successful']}\n"
            plain_output += f"Failed: {batch_result['batch_summary']['failed']}\n"
            plain_output += f"Total Processing Time: {batch_result['batch_summary']['total_processing_time']} seconds\n\n"
            
            if results:
                plain_output += "SUCCESSFUL RESULTS\n"
                plain_output += "-----------------\n\n"
                for i, result in enumerate(batch_result["results"], 1):
                    plain_output += f"{i}. {os.path.basename(result['file_path'])}\n"
                    plain_output += "=" * (len(os.path.basename(result['file_path'])) + 3) + "\n\n"
                    if isinstance(result["output"], str):
                        plain_output += result["output"] + "\n\n"
                    else:
                        plain_output += f"Title: {result['output'].get('title', 'Summary')}\n\n"
                        plain_output += f"{result['output'].get('summary', '')}\n\n"
                        
                        if "key_points" in result["output"]:
                            plain_output += "Key Points:\n"
                            for j, point in enumerate(result["output"]["key_points"], 1):
                                plain_output += f"{j}. {point}\n"
                            plain_output += "\n"
            
            if errors:
                plain_output += "ERRORS\n"
                plain_output += "------\n\n"
                for i, error in enumerate(batch_result["errors"], 1):
                    plain_output += f"{i}. {os.path.basename(error['file_path'])}\n"
                    plain_output += f"Error: {error['error']}\n"
                    plain_output += f"Recommendation: {error['recommendation']}\n\n"
            
            return plain_output
    else:
        # Return single file result
        if results:
            return results[0]["output"]
        else:
            error = errors[0]
            error_result = {
                "success": False,
                "error": error["error"],
                "error_type": error["error_type"],
                "file_path": error["file_path"],
                "recommendation": get_error_recommendation(error["error_type"], error["error"])
            }
            
            if output_format == "json":
                return json.dumps(error_result, indent=2)
            elif output_format == "markdown":
                md_output = f"# Error Processing {os.path.basename(error['file_path'])}\n\n"
                md_output += f"**Error:** {error['error']}\n\n"
                md_output += f"**Recommendation:** {error['recommendation']}\n"
                return md_output
            else:  # plain
                plain_output = f"ERROR PROCESSING {os.path.basename(error['file_path'])}\n"
                plain_output += "=" * (len(os.path.basename(error['file_path'])) + 16) + "\n\n"
                plain_output += f"Error: {error['error']}\n\n"
                plain_output += f"Recommendation: {error['recommendation']}\n"
                return plain_output


def get_error_recommendation(error_type: str, error_message: str) -> str:
    """Get recommendation based on error type"""
    if "UnsupportedOperation" in error_type or "Unsupported file type" in error_message:
        return "The file format is not supported. Please convert to one of the supported formats: PDF, DOCX, TXT, or MD."
    elif "FileNotFoundError" in error_type:
        return "The specified file was not found. Please check the file path and try again."
    elif "PermissionError" in error_type:
        return "Permission denied when accessing the file. Please check file permissions."
    elif "ValueError" in error_type and "encrypted" in error_message.lower():
        return "The file appears to be encrypted or password-protected. Please provide an unprotected version of the file."
    elif "decrypt" in error_message.lower() or "password" in error_message.lower():
        return "The file is password-protected. Please provide an unprotected version of the file."
    elif "RuntimeError" in error_type and "summary" in error_message.lower():
        return "Failed to generate summary. This might be due to API limits, network issues, or content that is too complex. Try with a smaller file or different content."
    elif "memory" in error_message.lower():
        return "The file is too large to process. Try splitting it into smaller files or use a file that is under 10MB."
    else:
        return "An unexpected error occurred. Please try again with a different file or contact support if the issue persists."