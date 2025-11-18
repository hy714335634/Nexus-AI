#!/usr/bin/env python3
"""
DOCX Document Parser Tool

This module provides tools for parsing DOCX documents, extracting metadata,
caching document content by sections, and retrieving document content by page or section.
"""

import os
import json
import uuid
import base64
import shutil
import logging
from typing import Dict, Any, List, Optional, Tuple, Union
from pathlib import Path
import math

# Third-party imports
from strands import tool
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docx.document import Document as DocumentType
from docx.text.paragraph import Paragraph
from docx.table import Table
from PIL import Image
from io import BytesIO

# Configure logging
logger = logging.getLogger(__name__)

# Constants
DEFAULT_CACHE_DIR = ".cache/docx_parser"
CHARS_PER_PAGE = 3000  # Approximate number of characters per page
STYLES_HEADING = ['Heading', 'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5', 'Heading 6', 'Title']

class DocxParsingError(Exception):
    """Exception raised for errors during DOCX parsing."""
    pass

def ensure_cache_dir(job_id: str) -> Tuple[str, str, str]:
    """
    Ensures that cache directories for a job exist.
    
    Args:
        job_id: The job ID
        
    Returns:
        Tuple containing paths to job directory, sections directory, and images directory
    """
    job_dir = os.path.join(DEFAULT_CACHE_DIR, job_id)
    sections_dir = os.path.join(job_dir, "sections")
    images_dir = os.path.join(job_dir, "images")
    
    for directory in [job_dir, sections_dir, images_dir]:
        os.makedirs(directory, exist_ok=True)
    
    return job_dir, sections_dir, images_dir

def is_heading(paragraph: Paragraph) -> bool:
    """
    Determines if a paragraph is a heading.
    
    Args:
        paragraph: The paragraph to check
        
    Returns:
        True if the paragraph is a heading, False otherwise
    """
    if not paragraph.text.strip():
        return False
    
    if paragraph.style.name in STYLES_HEADING:
        return True
    
    # Also check for custom heading styles
    style_name = paragraph.style.name.lower()
    return any(heading.lower() in style_name for heading in STYLES_HEADING)

def get_heading_level(paragraph: Paragraph) -> int:
    """
    Gets the heading level of a paragraph.
    
    Args:
        paragraph: The paragraph to check
        
    Returns:
        The heading level (1-6), or 0 if not a heading
    """
    if not is_heading(paragraph):
        return 0
    
    style_name = paragraph.style.name.lower()
    
    # Check for numbered heading styles
    for i in range(1, 7):
        if f"heading {i}" in style_name:
            return i
    
    # Default levels for other heading styles
    if "title" in style_name:
        return 1
    if "heading" in style_name:
        return 2
    
    return 1  # Default level for any other heading style

def extract_images(doc: DocumentType, images_dir: str) -> List[Dict[str, Any]]:
    """
    Extracts images from the document and saves them to the images directory.
    
    Args:
        doc: The document object
        images_dir: The directory to save images to
        
    Returns:
        List of dictionaries containing image metadata
    """
    images = []
    image_id = 1
    
    # Process all document parts
    for rel in doc.part.rels.values():
        # Check if the relationship is an image
        if "image" in rel.reltype:
            try:
                # Get image data
                image_data = rel.target_part.blob
                
                # Use PIL to get image info
                with BytesIO(image_data) as image_stream:
                    img = Image.open(image_stream)
                    format_ext = img.format.lower()
                    width, height = img.size
                
                # Generate image ID and filename
                image_name = f"image_{image_id:03d}.{format_ext}"
                image_path = os.path.join(images_dir, image_name)
                
                # Save the image
                with open(image_path, "wb") as f:
                    f.write(image_data)
                
                # Store image metadata
                images.append({
                    "id": f"image_{image_id:03d}",
                    "filename": image_name,
                    "path": image_path,
                    "format": format_ext,
                    "width": width,
                    "height": height,
                    "size_bytes": len(image_data)
                })
                
                image_id += 1
            except Exception as e:
                logger.warning(f"Failed to process image: {e}")
    
    return images

def parse_document_structure(doc: DocumentType) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]], int]:
    """
    Parses the document structure to extract sections and content.
    
    Args:
        doc: The document object
        
    Returns:
        Tuple containing sections list, content list, and total character count
    """
    sections = []
    content = []
    char_count = 0
    current_section = None
    section_id = 0
    
    # Process each paragraph in the document
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        char_count += len(text)
        
        # Check if this paragraph is a heading
        if is_heading(para):
            # If we already have a section, add it to the list
            if current_section:
                sections.append(current_section)
            
            # Start a new section
            section_id += 1
            level = get_heading_level(para)
            current_section = {
                "id": f"section_{section_id}",
                "title": text,
                "level": level,
                "start_index": i,
                "content_indices": []
            }
        
        # If we have a current section, add this paragraph index to it
        if current_section:
            current_section["content_indices"].append(i)
        
        # Add paragraph to content list
        content.append({
            "index": i,
            "type": "paragraph",
            "text": text,
            "is_heading": is_heading(para),
            "heading_level": get_heading_level(para) if is_heading(para) else 0
        })
    
    # Add the last section if there is one
    if current_section:
        sections.append(current_section)
    
    # If no sections were found, create a default one
    if not sections:
        sections.append({
            "id": "section_1",
            "title": "Document",
            "level": 0,
            "start_index": 0,
            "content_indices": list(range(len(doc.paragraphs)))
        })
    
    return sections, content, char_count

def estimate_pages(char_count: int, image_count: int) -> int:
    """
    Estimates the number of pages in the document.
    
    Args:
        char_count: Total character count
        image_count: Total image count
        
    Returns:
        Estimated number of pages
    """
    # Estimate pages based on character count
    char_pages = math.ceil(char_count / CHARS_PER_PAGE)
    
    # Add pages for images (assuming larger images take more space)
    image_pages = math.ceil(image_count / 2)
    
    return max(1, char_pages + image_pages)

def map_content_to_pages(content: List[Dict[str, Any]], images: List[Dict[str, Any]], page_count: int) -> Dict[int, List[int]]:
    """
    Maps content indices to page numbers.
    
    Args:
        content: Content list
        images: Images list
        page_count: Total page count
        
    Returns:
        Dictionary mapping page numbers to content indices
    """
    # Calculate total content items
    total_items = len(content) + len(images)
    
    # Calculate items per page
    items_per_page = math.ceil(total_items / page_count) if page_count > 0 else total_items
    
    # Create mapping
    page_map = {}
    current_page = 1
    current_items = 0
    
    # Map content indices to pages
    for i, _ in enumerate(content):
        if current_page not in page_map:
            page_map[current_page] = []
        
        page_map[current_page].append(i)
        current_items += 1
        
        if current_items >= items_per_page and current_page < page_count:
            current_page += 1
            current_items = 0
    
    return page_map

def save_metadata(job_dir: str, job_id: str, sections: List[Dict[str, Any]], 
                  content: List[Dict[str, Any]], images: List[Dict[str, Any]], 
                  page_map: Dict[int, List[int]], char_count: int, page_count: int) -> Dict[str, Any]:
    """
    Saves metadata to the job directory.
    
    Args:
        job_dir: Job directory path
        job_id: Job ID
        sections: Sections list
        content: Content list
        images: Images list
        page_map: Page mapping
        char_count: Total character count
        page_count: Total page count
        
    Returns:
        Metadata dictionary
    """
    metadata = {
        "job_id": job_id,
        "document": {
            "page_count": page_count,
            "char_count": char_count,
            "section_count": len(sections),
            "image_count": len(images)
        },
        "sections": [{"id": section["id"], "title": section["title"], "level": section["level"]} for section in sections],
        "images": [{"id": image["id"], "filename": image["filename"]} for image in images],
        "page_mapping": {str(page): indices for page, indices in page_map.items()}
    }
    
    # Save metadata to file
    with open(os.path.join(job_dir, "metadata.json"), "w", encoding="utf-8") as f:
        json.dump(metadata, f, ensure_ascii=False, indent=2)
    
    return metadata

def save_sections(sections_dir: str, sections: List[Dict[str, Any]], content: List[Dict[str, Any]]) -> None:
    """
    Saves section content to section files.
    
    Args:
        sections_dir: Sections directory path
        sections: Sections list
        content: Content list
    """
    for section in sections:
        section_content = [content[i] for i in section["content_indices"]]
        
        # Save section content to file
        with open(os.path.join(sections_dir, f"{section['id']}.json"), "w", encoding="utf-8") as f:
            json.dump({
                "id": section["id"],
                "title": section["title"],
                "level": section["level"],
                "content": section_content
            }, f, ensure_ascii=False, indent=2)

@tool
def parse_document(doc_path: str, job_id: Optional[str] = None) -> str:
    """
    Parses a DOCX document and extracts metadata.
    
    Args:
        doc_path: Path to the DOCX document
        job_id: Optional job ID, if not provided a new one will be generated
        
    Returns:
        JSON string containing job_id, page count, character count, section information, etc.
    """
    try:
        # Validate file existence
        if not os.path.exists(doc_path):
            return json.dumps({
                "error": "File not found",
                "message": f"Document not found at path: {doc_path}"
            }, ensure_ascii=False, indent=2)
        
        # Validate file extension
        if not doc_path.lower().endswith(".docx"):
            return json.dumps({
                "error": "Invalid file format",
                "message": "Only DOCX files are supported"
            }, ensure_ascii=False, indent=2)
        
        # Generate job ID if not provided
        if not job_id:
            job_id = str(uuid.uuid4())
        
        # Ensure cache directories exist
        job_dir, sections_dir, images_dir = ensure_cache_dir(job_id)
        
        # Parse document
        try:
            doc = Document(doc_path)
        except PackageNotFoundError:
            return json.dumps({
                "error": "Invalid DOCX file",
                "message": "The file is not a valid DOCX document"
            }, ensure_ascii=False, indent=2)
        
        # Extract images
        images = extract_images(doc, images_dir)
        
        # Parse document structure
        sections, content, char_count = parse_document_structure(doc)
        
        # Estimate page count
        page_count = estimate_pages(char_count, len(images))
        
        # Map content to pages
        page_map = map_content_to_pages(content, images, page_count)
        
        # Save metadata
        metadata = save_metadata(
            job_dir, job_id, sections, content, images, 
            page_map, char_count, page_count
        )
        
        # Save sections
        save_sections(sections_dir, sections, content)
        
        # Return metadata
        return json.dumps({
            "status": "success",
            "job_id": job_id,
            "document": metadata["document"],
            "sections": metadata["sections"],
            "images": metadata["images"]
        }, ensure_ascii=False, indent=2)
        
    except Exception as e:
        logger.error(f"Error parsing document: {e}")
        return json.dumps({
            "error": "Document parsing failed",
            "message": str(e)
        }, ensure_ascii=False, indent=2)

@tool
def get_page_content(job_id: str, page_num: int) -> str:
    """
    Gets content for a specific page.
    
    Args:
        job_id: Job ID
        page_num: Page number
        
    Returns:
        JSON string containing the page content
    """
    try:
        # Validate job exists
        job_dir = os.path.join(DEFAULT_CACHE_DIR, job_id)
        if not os.path.exists(job_dir):
            return json.dumps({
                "error": "Job not found",
                "message": f"No parsed document found for job ID: {job_id}"
            }, ensure_ascii=False, indent=2)
        
        # Load metadata
        metadata_path = os.path.join(job_dir, "metadata.json")
        with open(metadata_path, "r", encoding="utf-8") as f:
            metadata = json.load(f)
        
        # Validate page number
        if page_num < 1 or page_num > metadata["document"]["page_count"]:
            return json.dumps({
                "error": "Invalid page number",
                "message": f"Page number must be between 1 and {metadata['document']['page_count']}"
            }, ensure_ascii=False, indent=2)
        
        # Get content indices for the page
        content_indices = metadata["page_mapping"].get(str(page_num), [])
        
        # Load all sections to collect content
        sections_dir = os.path.join(job_dir, "sections")
        all_content = []
        
        # We need to collect all content from all sections first
        for section_info in metadata["sections"]:
            section_path = os.path.join(sections_dir, f"{section_info['id']}.json")
            with open(section_path, "r", encoding="utf-8") as f:
                section_data = json.load(f)
                for item in section_data["content"]:
                    all_content.append(item)
        
        # Get content for the page
        page_content = [all_content[i] for i in content_indices if i < len(all_content)]
        
        # Get section information for content on this page
        page_sections = []
        for item in page_content:
            if item["is_heading"]:
                for section in metadata["sections"]:
                    if section["title"] == item["text"]:
                        page_sections.append(section)
                        break
        
        return json.dumps({
            "status": "success",
            "job_id": job_id,
            "page": page_num,
            "content": page_content,
            "sections": page_sections
        }, ensure_ascii=False, indent=2)
        
    except Exception as e:
        logger.error(f"Error retrieving page content: {e}")
        return json.dumps({
            "error": "Failed to retrieve page content",
            "message": str(e)
        }, ensure_ascii=False, indent=2)

@tool
def get_section_content(job_id: str, section_title: str) -> str:
    """
    Gets content for a specific section.
    
    Args:
        job_id: Job ID
        section_title: Section title
        
    Returns:
        JSON string containing the section content
    """
    try:
        # Validate job exists
        job_dir = os.path.join(DEFAULT_CACHE_DIR, job_id)
        if not os.path.exists(job_dir):
            return json.dumps({
                "error": "Job not found",
                "message": f"No parsed document found for job ID: {job_id}"
            }, ensure_ascii=False, indent=2)
        
        # Load metadata
        metadata_path = os.path.join(job_dir, "metadata.json")
        with open(metadata_path, "r", encoding="utf-8") as f:
            metadata = json.load(f)
        
        # Find section by title
        section_id = None
        for section in metadata["sections"]:
            if section["title"] == section_title:
                section_id = section["id"]
                break
        
        if not section_id:
            return json.dumps({
                "error": "Section not found",
                "message": f"No section found with title: {section_title}"
            }, ensure_ascii=False, indent=2)
        
        # Load section content
        section_path = os.path.join(job_dir, "sections", f"{section_id}.json")
        with open(section_path, "r", encoding="utf-8") as f:
            section_data = json.load(f)
        
        # Find images that might be in this section
        # This is an approximation - we're looking at images between this section and the next
        sections_sorted = sorted(metadata["sections"], key=lambda s: int(s["id"].split("_")[1]))
        current_index = next((i for i, s in enumerate(sections_sorted) if s["id"] == section_id), -1)
        
        # Get images in this section (if any)
        section_images = []
        if "images" in metadata and metadata["images"]:
            section_images = metadata["images"]
        
        return json.dumps({
            "status": "success",
            "job_id": job_id,
            "section": {
                "id": section_data["id"],
                "title": section_data["title"],
                "level": section_data["level"]
            },
            "content": section_data["content"],
            "images": section_images
        }, ensure_ascii=False, indent=2)
        
    except Exception as e:
        logger.error(f"Error retrieving section content: {e}")
        return json.dumps({
            "error": "Failed to retrieve section content",
            "message": str(e)
        }, ensure_ascii=False, indent=2)

@tool
def get_image(job_id: str, image_id: str, return_type: str = "path") -> str:
    """
    Gets an image from the document.
    
    Args:
        job_id: Job ID
        image_id: Image ID
        return_type: Return type, either 'base64' or 'path'
        
    Returns:
        JSON string containing the image data
    """
    try:
        # Validate job exists
        job_dir = os.path.join(DEFAULT_CACHE_DIR, job_id)
        if not os.path.exists(job_dir):
            return json.dumps({
                "error": "Job not found",
                "message": f"No parsed document found for job ID: {job_id}"
            }, ensure_ascii=False, indent=2)
        
        # Load metadata
        metadata_path = os.path.join(job_dir, "metadata.json")
        with open(metadata_path, "r", encoding="utf-8") as f:
            metadata = json.load(f)
        
        # Find image
        image_info = None
        for image in metadata["images"]:
            if image["id"] == image_id:
                image_info = image
                break
        
        if not image_info:
            return json.dumps({
                "error": "Image not found",
                "message": f"No image found with ID: {image_id}"
            }, ensure_ascii=False, indent=2)
        
        # Get image path
        image_path = os.path.join(job_dir, "images", image_info["filename"])
        
        if not os.path.exists(image_path):
            return json.dumps({
                "error": "Image file not found",
                "message": f"Image file not found: {image_info['filename']}"
            }, ensure_ascii=False, indent=2)
        
        # Return based on return type
        if return_type.lower() == "base64":
            # Read image and convert to base64
            with open(image_path, "rb") as f:
                image_data = base64.b64encode(f.read()).decode("utf-8")
            
            return json.dumps({
                "status": "success",
                "job_id": job_id,
                "image_id": image_id,
                "return_type": "base64",
                "data": image_data,
                "format": os.path.splitext(image_info["filename"])[1][1:]
            }, ensure_ascii=False, indent=2)
        else:
            # Return path
            return json.dumps({
                "status": "success",
                "job_id": job_id,
                "image_id": image_id,
                "return_type": "path",
                "path": image_path,
                "format": os.path.splitext(image_info["filename"])[1][1:]
            }, ensure_ascii=False, indent=2)
        
    except Exception as e:
        logger.error(f"Error retrieving image: {e}")
        return json.dumps({
            "error": "Failed to retrieve image",
            "message": str(e)
        }, ensure_ascii=False, indent=2)

@tool
def clean_cache(job_id: str) -> str:
    """
    Cleans cache for a specific job.
    
    Args:
        job_id: Job ID
        
    Returns:
        JSON string containing the result
    """
    try:
        # Validate job exists
        job_dir = os.path.join(DEFAULT_CACHE_DIR, job_id)
        if not os.path.exists(job_dir):
            return json.dumps({
                "error": "Job not found",
                "message": f"No parsed document found for job ID: {job_id}"
            }, ensure_ascii=False, indent=2)
        
        # Remove job directory
        shutil.rmtree(job_dir)
        
        return json.dumps({
            "status": "success",
            "message": f"Cache cleaned for job ID: {job_id}"
        }, ensure_ascii=False, indent=2)
        
    except Exception as e:
        logger.error(f"Error cleaning cache: {e}")
        return json.dumps({
            "error": "Failed to clean cache",
            "message": str(e)
        }, ensure_ascii=False, indent=2)