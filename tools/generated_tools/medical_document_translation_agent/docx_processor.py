#!/usr/bin/env python3
"""
docx_processor.py - Microsoft Word文档处理工具

该工具提供了读取、解析、修改和生成Microsoft Word文档(.docx)的功能。
支持保持原文档的排版、布局、格式和样式。
"""

import os
import json
import tempfile
from typing import Dict, List, Any, Optional, Union, BinaryIO, Tuple
from io import BytesIO
import base64

try:
    import docx
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.table import Table
    from docx.oxml.ns import qn
except ImportError:
    raise ImportError("请安装python-docx库: pip install python-docx")

from strands import tool


@tool
def read_docx(file_path: str, return_format: str = "json", extract_images: bool = False) -> str:
    """
    读取和解析.docx文件，提取文档结构和内容
    
    Args:
        file_path (str): Word文档的路径
        return_format (str): 返回格式，可选值: "json", "text"
        extract_images (bool): 是否提取图片（以base64编码返回）
        
    Returns:
        str: 文档内容的JSON或文本表示
    """
    try:
        if not os.path.exists(file_path):
            return json.dumps({"error": f"文件不存在: {file_path}"})
            
        if not file_path.lower().endswith('.docx'):
            return json.dumps({"error": "仅支持.docx格式文件"})
        
        doc = Document(file_path)
        
        if return_format.lower() == "text":
            # 简单文本提取
            text_content = []
            for para in doc.paragraphs:
                text_content.append(para.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text)
                    text_content.append(" | ".join(row_text))
            
            return "\n".join(text_content)
        
        else:  # JSON格式
            document_data = {
                "metadata": {
                    "title": os.path.basename(file_path),
                    "core_properties": _extract_core_properties(doc),
                },
                "content": {
                    "paragraphs": [],
                    "tables": [],
                    "sections": [],
                    "headers": [],
                    "footers": [],
                }
            }
            
            # 提取段落
            for i, para in enumerate(doc.paragraphs):
                paragraph_data = {
                    "id": f"p{i}",
                    "text": para.text,
                    "style": para.style.name if para.style else "Normal",
                    "alignment": str(para.alignment) if para.alignment else "LEFT",
                    "runs": []
                }
                
                # 提取格式化文本块
                for j, run in enumerate(para.runs):
                    run_data = {
                        "id": f"p{i}r{j}",
                        "text": run.text,
                        "bold": run.bold,
                        "italic": run.italic,
                        "underline": run.underline,
                        "font": run.font.name if run.font.name else None,
                        "size": str(run.font.size) if run.font.size else None,
                        "color": str(run.font.color.rgb) if run.font.color and run.font.color.rgb else None,
                    }
                    paragraph_data["runs"].append(run_data)
                
                document_data["content"]["paragraphs"].append(paragraph_data)
            
            # 提取表格
            for i, table in enumerate(doc.tables):
                table_data = {
                    "id": f"t{i}",
                    "rows": len(table.rows),
                    "columns": len(table.columns),
                    "cells": []
                }
                
                for r, row in enumerate(table.rows):
                    for c, cell in enumerate(row.cells):
                        cell_data = {
                            "row": r,
                            "column": c,
                            "text": cell.text,
                            "paragraphs": []
                        }
                        
                        # 提取单元格中的段落
                        for p, para in enumerate(cell.paragraphs):
                            para_data = {
                                "id": f"t{i}r{r}c{c}p{p}",
                                "text": para.text,
                                "runs": []
                            }
                            
                            for j, run in enumerate(para.runs):
                                run_data = {
                                    "id": f"t{i}r{r}c{c}p{p}r{j}",
                                    "text": run.text,
                                    "bold": run.bold,
                                    "italic": run.italic,
                                    "underline": run.underline
                                }
                                para_data["runs"].append(run_data)
                                
                            cell_data["paragraphs"].append(para_data)
                            
                        table_data["cells"].append(cell_data)
                
                document_data["content"]["tables"].append(table_data)
            
            # 提取页眉页脚
            for i, section in enumerate(doc.sections):
                section_data = {
                    "id": f"s{i}",
                    "page_height": str(section.page_height),
                    "page_width": str(section.page_width),
                    "left_margin": str(section.left_margin),
                    "right_margin": str(section.right_margin),
                    "top_margin": str(section.top_margin),
                    "bottom_margin": str(section.bottom_margin),
                    "header_distance": str(section.header_distance),
                    "footer_distance": str(section.footer_distance),
                    "orientation": "PORTRAIT" if section.orientation == 0 else "LANDSCAPE"
                }
                document_data["content"]["sections"].append(section_data)
                
                # 页眉
                if section.header:
                    header_data = {
                        "section_id": f"s{i}",
                        "paragraphs": []
                    }
                    
                    for p, para in enumerate(section.header.paragraphs):
                        if para.text.strip():  # 只添加非空段落
                            para_data = {
                                "id": f"s{i}h{p}",
                                "text": para.text
                            }
                            header_data["paragraphs"].append(para_data)
                            
                    if header_data["paragraphs"]:
                        document_data["content"]["headers"].append(header_data)
                
                # 页脚
                if section.footer:
                    footer_data = {
                        "section_id": f"s{i}",
                        "paragraphs": []
                    }
                    
                    for p, para in enumerate(section.footer.paragraphs):
                        if para.text.strip():  # 只添加非空段落
                            para_data = {
                                "id": f"s{i}f{p}",
                                "text": para.text
                            }
                            footer_data["paragraphs"].append(para_data)
                            
                    if footer_data["paragraphs"]:
                        document_data["content"]["footers"].append(footer_data)
            
            # 提取图片（如果需要）
            if extract_images:
                document_data["content"]["images"] = _extract_images(doc)
            
            # 提取样式信息
            document_data["styles"] = _extract_styles(doc)
            
            return json.dumps(document_data, ensure_ascii=False, indent=2)
            
    except Exception as e:
        return json.dumps({"error": f"处理Word文档时出错: {str(e)}"})


@tool
def create_docx(content: Dict[str, Any], output_path: str) -> str:
    """
    根据提供的内容创建Word文档
    
    Args:
        content (Dict[str, Any]): 文档内容的结构化表示（JSON格式）
        output_path (str): 输出文档的保存路径
        
    Returns:
        str: 操作结果信息
    """
    try:
        doc = Document()
        
        # 应用元数据
        if "metadata" in content:
            metadata = content["metadata"]
            if "core_properties" in metadata:
                props = metadata["core_properties"]
                if "title" in props:
                    doc.core_properties.title = props["title"]
                if "author" in props:
                    doc.core_properties.author = props["author"]
                if "subject" in props:
                    doc.core_properties.subject = props["subject"]
                if "keywords" in props:
                    doc.core_properties.keywords = props["keywords"]
                if "category" in props:
                    doc.core_properties.category = props["category"]
        
        # 应用样式
        if "styles" in content:
            _apply_styles(doc, content["styles"])
        
        # 创建文档内容
        if "content" in content:
            doc_content = content["content"]
            
            # 添加段落
            if "paragraphs" in doc_content:
                for para_data in doc_content["paragraphs"]:
                    para = doc.add_paragraph()
                    
                    # 设置段落样式
                    if "style" in para_data:
                        para.style = para_data["style"]
                    
                    # 设置对齐方式
                    if "alignment" in para_data:
                        alignment_map = {
                            "LEFT": WD_PARAGRAPH_ALIGNMENT.LEFT,
                            "CENTER": WD_PARAGRAPH_ALIGNMENT.CENTER,
                            "RIGHT": WD_PARAGRAPH_ALIGNMENT.RIGHT,
                            "JUSTIFY": WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        }
                        if para_data["alignment"] in alignment_map:
                            para.alignment = alignment_map[para_data["alignment"]]
                    
                    # 添加文本运行
                    if "runs" in para_data and para_data["runs"]:
                        for run_data in para_data["runs"]:
                            run = para.add_run(run_data.get("text", ""))
                            
                            # 应用文本格式
                            run.bold = run_data.get("bold", False)
                            run.italic = run_data.get("italic", False)
                            run.underline = run_data.get("underline", False)
                            
                            # 应用字体
                            if "font" in run_data and run_data["font"]:
                                run.font.name = run_data["font"]
                            
                            # 应用字体大小
                            if "size" in run_data and run_data["size"]:
                                try:
                                    size_value = run_data["size"]
                                    if isinstance(size_value, str) and "Pt" in size_value:
                                        size_pt = float(size_value.replace("Pt", "").strip())
                                        run.font.size = Pt(size_pt)
                                    else:
                                        run.font.size = Pt(float(size_value))
                                except:
                                    pass  # 忽略无效的字体大小
                            
                            # 应用颜色
                            if "color" in run_data and run_data["color"]:
                                try:
                                    color_str = run_data["color"]
                                    if color_str.startswith("RGB"):
                                        rgb_values = color_str.replace("RGB", "").replace("(", "").replace(")", "").split(",")
                                        if len(rgb_values) == 3:
                                            r, g, b = map(int, rgb_values)
                                            run.font.color.rgb = RGBColor(r, g, b)
                                except:
                                    pass  # 忽略无效的颜色
                    else:
                        # 如果没有runs，直接添加文本
                        para.text = para_data.get("text", "")
            
            # 添加表格
            if "tables" in doc_content:
                for table_data in doc_content["tables"]:
                    rows = table_data.get("rows", 1)
                    cols = table_data.get("columns", 1)
                    
                    table = doc.add_table(rows=rows, cols=cols)
                    table.style = 'Table Grid'
                    
                    # 填充单元格
                    if "cells" in table_data:
                        for cell_data in table_data["cells"]:
                            row_idx = cell_data.get("row", 0)
                            col_idx = cell_data.get("column", 0)
                            
                            if row_idx < rows and col_idx < cols:
                                cell = table.cell(row_idx, col_idx)
                                
                                # 清除默认段落
                                cell.paragraphs[0].clear()
                                
                                # 添加段落
                                if "paragraphs" in cell_data:
                                    for para_data in cell_data["paragraphs"]:
                                        cell_para = cell.add_paragraph(para_data.get("text", ""))
                                        
                                        # 添加格式化文本
                                        if "runs" in para_data:
                                            cell_para.clear()  # 清除默认文本
                                            for run_data in para_data["runs"]:
                                                cell_run = cell_para.add_run(run_data.get("text", ""))
                                                cell_run.bold = run_data.get("bold", False)
                                                cell_run.italic = run_data.get("italic", False)
                                                cell_run.underline = run_data.get("underline", False)
                                else:
                                    # 如果没有段落数据，直接设置文本
                                    cell.text = cell_data.get("text", "")
            
            # 设置页面布局
            if "sections" in doc_content:
                # 确保至少有一个节
                if not doc.sections:
                    doc.add_section()
                
                for i, section_data in enumerate(doc_content["sections"]):
                    # 如果需要更多节，添加新节
                    if i > 0 and i < len(doc.sections):
                        doc.add_section()
                    
                    # 获取当前节
                    section = doc.sections[i] if i < len(doc.sections) else doc.sections[-1]
                    
                    # 设置页面大小和边距
                    if "page_width" in section_data and "page_height" in section_data:
                        try:
                            section.page_width = _parse_measurement(section_data["page_width"])
                            section.page_height = _parse_measurement(section_data["page_height"])
                        except:
                            pass  # 忽略无效的页面尺寸
                    
                    if "left_margin" in section_data:
                        section.left_margin = _parse_measurement(section_data["left_margin"])
                    if "right_margin" in section_data:
                        section.right_margin = _parse_measurement(section_data["right_margin"])
                    if "top_margin" in section_data:
                        section.top_margin = _parse_measurement(section_data["top_margin"])
                    if "bottom_margin" in section_data:
                        section.bottom_margin = _parse_measurement(section_data["bottom_margin"])
                    
                    # 设置页面方向
                    if "orientation" in section_data:
                        from docx.enum.section import WD_ORIENTATION
                        orientation = section_data["orientation"]
                        if orientation == "LANDSCAPE":
                            section.orientation = WD_ORIENTATION.LANDSCAPE
                        else:
                            section.orientation = WD_ORIENTATION.PORTRAIT
            
            # 添加页眉
            if "headers" in doc_content:
                for header_data in doc_content["headers"]:
                    section_id = header_data.get("section_id", "s0")
                    section_idx = int(section_id.replace("s", ""))
                    
                    if section_idx < len(doc.sections):
                        section = doc.sections[section_idx]
                        header = section.header
                        
                        # 清除默认页眉
                        for para in header.paragraphs:
                            para.clear()
                        
                        # 添加页眉内容
                        if "paragraphs" in header_data:
                            for para_data in header_data["paragraphs"]:
                                header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                                header_para.text = para_data.get("text", "")
            
            # 添加页脚
            if "footers" in doc_content:
                for footer_data in doc_content["footers"]:
                    section_id = footer_data.get("section_id", "s0")
                    section_idx = int(section_id.replace("s", ""))
                    
                    if section_idx < len(doc.sections):
                        section = doc.sections[section_idx]
                        footer = section.footer
                        
                        # 清除默认页脚
                        for para in footer.paragraphs:
                            para.clear()
                        
                        # 添加页脚内容
                        if "paragraphs" in footer_data:
                            for para_data in footer_data["paragraphs"]:
                                footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                                footer_para.text = para_data.get("text", "")
        
        # 保存文档
        doc.save(output_path)
        return json.dumps({
            "success": True,
            "message": f"文档已成功保存到 {output_path}",
            "file_path": output_path
        }, ensure_ascii=False)
        
    except Exception as e:
        return json.dumps({
            "success": False,
            "error": f"创建Word文档时出错: {str(e)}"
        }, ensure_ascii=False)


@tool
def modify_docx(file_path: str, modifications: Dict[str, Any], output_path: str) -> str:
    """
    修改现有Word文档
    
    Args:
        file_path (str): 要修改的Word文档路径
        modifications (Dict[str, Any]): 要应用的修改
        output_path (str): 修改后文档的保存路径
        
    Returns:
        str: 操作结果信息
    """
    try:
        if not os.path.exists(file_path):
            return json.dumps({"error": f"文件不存在: {file_path}"})
            
        if not file_path.lower().endswith('.docx'):
            return json.dumps({"error": "仅支持.docx格式文件"})
        
        # 加载文档
        doc = Document(file_path)
        
        # 应用修改
        if "replace_text" in modifications:
            for replace_item in modifications["replace_text"]:
                old_text = replace_item.get("old", "")
                new_text = replace_item.get("new", "")
                
                # 替换段落中的文本
                for para in doc.paragraphs:
                    if old_text in para.text:
                        for run in para.runs:
                            run.text = run.text.replace(old_text, new_text)
                
                # 替换表格中的文本
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                if old_text in para.text:
                                    for run in para.runs:
                                        run.text = run.text.replace(old_text, new_text)
        
        # 添加新段落
        if "add_paragraphs" in modifications:
            for para_data in modifications["add_paragraphs"]:
                text = para_data.get("text", "")
                style = para_data.get("style", "Normal")
                position = para_data.get("position", "end")  # "end" 或段落索引
                
                if position == "end":
                    para = doc.add_paragraph(text)
                    para.style = style
                else:
                    try:
                        idx = int(position)
                        if 0 <= idx <= len(doc.paragraphs):
                            # 在指定位置插入段落
                            p = doc.paragraphs[idx]
                            new_para = p._element.addnext(docx.oxml.shared.OxmlElement('w:p'))
                            new_para = docx.text.paragraph.Paragraph(new_para, p._parent)
                            new_para.text = text
                            new_para.style = style
                    except:
                        # 如果位置无效，添加到末尾
                        para = doc.add_paragraph(text)
                        para.style = style
        
        # 添加新表格
        if "add_tables" in modifications:
            for table_data in modifications["add_tables"]:
                rows = table_data.get("rows", 1)
                cols = table_data.get("columns", 1)
                data = table_data.get("data", [])
                
                table = doc.add_table(rows=rows, cols=cols)
                table.style = table_data.get("style", "Table Grid")
                
                # 填充表格数据
                for r, row_data in enumerate(data):
                    if r < rows:
                        for c, cell_text in enumerate(row_data):
                            if c < cols:
                                table.cell(r, c).text = str(cell_text)
        
        # 修改样式
        if "modify_styles" in modifications:
            for style_data in modifications["modify_styles"]:
                style_name = style_data.get("name", "")
                if style_name and style_name in doc.styles:
                    style = doc.styles[style_name]
                    
                    # 修改字体
                    if "font" in style_data:
                        font_data = style_data["font"]
                        if "name" in font_data:
                            style.font.name = font_data["name"]
                        if "size" in font_data:
                            try:
                                style.font.size = Pt(float(font_data["size"]))
                            except:
                                pass
                        if "bold" in font_data:
                            style.font.bold = font_data["bold"]
                        if "italic" in font_data:
                            style.font.italic = font_data["italic"]
        
        # 保存修改后的文档
        doc.save(output_path)
        
        return json.dumps({
            "success": True,
            "message": f"文档已成功修改并保存到 {output_path}",
            "file_path": output_path
        }, ensure_ascii=False)
        
    except Exception as e:
        return json.dumps({
            "success": False,
            "error": f"修改Word文档时出错: {str(e)}"
        }, ensure_ascii=False)


@tool
def merge_docx(file_paths: List[str], output_path: str) -> str:
    """
    合并多个Word文档
    
    Args:
        file_paths (List[str]): 要合并的Word文档路径列表
        output_path (str): 合并后文档的保存路径
        
    Returns:
        str: 操作结果信息
    """
    try:
        if not file_paths:
            return json.dumps({"error": "没有提供文件路径"})
        
        # 检查所有文件是否存在且为.docx格式
        for file_path in file_paths:
            if not os.path.exists(file_path):
                return json.dumps({"error": f"文件不存在: {file_path}"})
                
            if not file_path.lower().endswith('.docx'):
                return json.dumps({"error": f"仅支持.docx格式文件: {file_path}"})
        
        # 创建新文档作为合并结果
        merged_doc = Document()
        
        # 处理第一个文档
        first_doc = Document(file_paths[0])
        
        # 复制第一个文档的样式
        for style in first_doc.styles:
            if style.name not in merged_doc.styles:
                try:
                    merged_doc.styles.add_style(style.name, style.type)
                except:
                    pass  # 忽略内置样式
        
        # 复制第一个文档的内容
        for element in first_doc.element.body:
            merged_doc.element.body.append(element)
        
        # 处理其余文档
        for file_path in file_paths[1:]:
            doc = Document(file_path)
            
            # 添加分节符
            merged_doc.add_page_break()
            
            # 复制文档内容
            for element in doc.element.body:
                merged_doc.element.body.append(element)
        
        # 保存合并后的文档
        merged_doc.save(output_path)
        
        return json.dumps({
            "success": True,
            "message": f"已成功合并 {len(file_paths)} 个文档并保存到 {output_path}",
            "file_path": output_path
        }, ensure_ascii=False)
        
    except Exception as e:
        return json.dumps({
            "success": False,
            "error": f"合并Word文档时出错: {str(e)}"
        }, ensure_ascii=False)


@tool
def extract_docx_structure(file_path: str) -> str:
    """
    提取Word文档的结构信息（标题层次、目录等）
    
    Args:
        file_path (str): Word文档的路径
        
    Returns:
        str: JSON格式的文档结构信息
    """
    try:
        if not os.path.exists(file_path):
            return json.dumps({"error": f"文件不存在: {file_path}"})
            
        if not file_path.lower().endswith('.docx'):
            return json.dumps({"error": "仅支持.docx格式文件"})
        
        doc = Document(file_path)
        
        # 提取文档结构
        structure = {
            "title": os.path.basename(file_path),
            "headings": [],
            "toc": [],
            "sections": len(doc.sections),
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "styles_used": set()
        }
        
        # 提取标题层次
        current_level = 0
        heading_stack = []
        
        for i, para in enumerate(doc.paragraphs):
            # 记录使用的样式
            if para.style.name:
                structure["styles_used"].add(para.style.name)
            
            # 检查是否是标题样式
            if para.style.name.startswith('Heading'):
                try:
                    level = int(para.style.name.replace('Heading ', ''))
                    
                    heading = {
                        "level": level,
                        "text": para.text,
                        "index": i,
                        "children": []
                    }
                    
                    # 处理标题层次
                    if level == 1:
                        # 顶级标题
                        structure["headings"].append(heading)
                        heading_stack = [heading]
                        current_level = 1
                    elif level > current_level:
                        # 子标题
                        if heading_stack:
                            heading_stack[-1]["children"].append(heading)
                            heading_stack.append(heading)
                            current_level = level
                    elif level == current_level:
                        # 同级标题
                        if len(heading_stack) > 1:
                            heading_stack.pop()
                            heading_stack[-1]["children"].append(heading)
                            heading_stack.append(heading)
                    else:
                        # 返回上级标题
                        while current_level > level and len(heading_stack) > 1:
                            heading_stack.pop()
                            current_level -= 1
                        
                        if heading_stack:
                            if len(heading_stack) > 1:
                                heading_stack.pop()
                            
                            if level == 1:
                                structure["headings"].append(heading)
                            else:
                                heading_stack[-1]["children"].append(heading)
                            
                            heading_stack.append(heading)
                            current_level = level
                except:
                    pass  # 忽略无效的标题级别
            
            # 检查是否是目录条目
            if "TOC" in para.style.name:
                structure["toc"].append({
                    "text": para.text,
                    "index": i
                })
        
        # 将样式集合转换为列表
        structure["styles_used"] = list(structure["styles_used"])
        
        return json.dumps(structure, ensure_ascii=False, indent=2)
        
    except Exception as e:
        return json.dumps({
            "error": f"提取文档结构时出错: {str(e)}"
        }, ensure_ascii=False)


@tool
def create_docx_from_text(content: str, output_path: str, title: str = "翻译文档") -> str:
    """
    从文本内容创建Word文档（简化版本）
    
    Args:
        content (str): 文档内容文本
        output_path (str): 输出文档的保存路径
        title (str): 文档标题
        
    Returns:
        str: 操作结果信息
    """
    try:
        doc = Document()
        
        # 设置文档标题
        doc.core_properties.title = title
        
        # 将文本内容按段落分割
        paragraphs = content.split('\n')
        
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:  # 只添加非空段落
                # 检查是否是标题（以#开头）
                if para_text.startswith('#'):
                    # 计算标题级别
                    level = len(para_text) - len(para_text.lstrip('#'))
                    para_text = para_text.lstrip('# ').strip()
                    
                    # 添加标题段落
                    para = doc.add_paragraph(para_text)
                    if level == 1:
                        para.style = 'Heading 1'
                    elif level == 2:
                        para.style = 'Heading 2'
                    elif level == 3:
                        para.style = 'Heading 3'
                    else:
                        para.style = 'Heading 4'
                else:
                    # 添加普通段落
                    doc.add_paragraph(para_text)
        
        # 保存文档
        doc.save(output_path)
        return json.dumps({
            "success": True,
            "message": f"文档已成功保存到 {output_path}",
            "file_path": output_path
        }, ensure_ascii=False)
        
    except Exception as e:
        return json.dumps({
            "success": False,
            "error": f"创建Word文档时出错: {str(e)}"
        }, ensure_ascii=False)


@tool
def compare_docx(file_path1: str, file_path2: str) -> str:
    """
    比较两个Word文档的内容和格式差异
    
    Args:
        file_path1 (str): 第一个Word文档的路径
        file_path2 (str): 第二个Word文档的路径
        
    Returns:
        str: JSON格式的比较结果
    """
    try:
        # 检查文件是否存在
        for file_path in [file_path1, file_path2]:
            if not os.path.exists(file_path):
                return json.dumps({"error": f"文件不存在: {file_path}"})
                
            if not file_path.lower().endswith('.docx'):
                return json.dumps({"error": f"仅支持.docx格式文件: {file_path}"})
        
        # 加载两个文档
        doc1 = Document(file_path1)
        doc2 = Document(file_path2)
        
        # 比较结果
        comparison = {
            "document1": os.path.basename(file_path1),
            "document2": os.path.basename(file_path2),
            "paragraph_count": {
                "document1": len(doc1.paragraphs),
                "document2": len(doc2.paragraphs),
                "difference": len(doc2.paragraphs) - len(doc1.paragraphs)
            },
            "table_count": {
                "document1": len(doc1.tables),
                "document2": len(doc2.tables),
                "difference": len(doc2.tables) - len(doc1.tables)
            },
            "section_count": {
                "document1": len(doc1.sections),
                "document2": len(doc2.sections),
                "difference": len(doc2.sections) - len(doc1.sections)
            },
            "content_differences": [],
            "style_differences": []
        }
        
        # 比较段落内容
        min_paragraphs = min(len(doc1.paragraphs), len(doc2.paragraphs))
        for i in range(min_paragraphs):
            para1 = doc1.paragraphs[i]
            para2 = doc2.paragraphs[i]
            
            if para1.text != para2.text:
                comparison["content_differences"].append({
                    "paragraph_index": i,
                    "document1_text": para1.text,
                    "document2_text": para2.text
                })
            
            # 比较段落样式
            if para1.style.name != para2.style.name:
                comparison["style_differences"].append({
                    "paragraph_index": i,
                    "document1_style": para1.style.name,
                    "document2_style": para2.style.name
                })
        
        # 比较表格内容
        min_tables = min(len(doc1.tables), len(doc2.tables))
        table_differences = []
        
        for i in range(min_tables):
            table1 = doc1.tables[i]
            table2 = doc2.tables[i]
            
            # 比较表格尺寸
            if len(table1.rows) != len(table2.rows) or len(table1.columns) != len(table2.columns):
                table_differences.append({
                    "table_index": i,
                    "size_difference": {
                        "document1": f"{len(table1.rows)}x{len(table1.columns)}",
                        "document2": f"{len(table2.rows)}x{len(table2.columns)}"
                    }
                })
                continue
            
            # 比较单元格内容
            cell_differences = []
            for r in range(len(table1.rows)):
                for c in range(len(table1.columns)):
                    cell1 = table1.cell(r, c)
                    cell2 = table2.cell(r, c)
                    
                    if cell1.text != cell2.text:
                        cell_differences.append({
                            "row": r,
                            "column": c,
                            "document1_text": cell1.text,
                            "document2_text": cell2.text
                        })
            
            if cell_differences:
                table_differences.append({
                    "table_index": i,
                    "cell_differences": cell_differences
                })
        
        if table_differences:
            comparison["table_differences"] = table_differences
        
        # 计算总体差异度
        content_diff_count = len(comparison["content_differences"])
        style_diff_count = len(comparison["style_differences"])
        table_diff_count = len(table_differences)
        
        total_elements1 = len(doc1.paragraphs) + sum(len(t.rows) * len(t.columns) for t in doc1.tables)
        total_elements2 = len(doc2.paragraphs) + sum(len(t.rows) * len(t.columns) for t in doc2.tables)
        total_elements = max(total_elements1, total_elements2)
        
        total_diff_count = content_diff_count + style_diff_count + table_diff_count
        diff_percentage = (total_diff_count / total_elements * 100) if total_elements > 0 else 0
        
        comparison["difference_summary"] = {
            "content_differences": content_diff_count,
            "style_differences": style_diff_count,
            "table_differences": table_diff_count,
            "total_differences": total_diff_count,
            "difference_percentage": round(diff_percentage, 2)
        }
        
        return json.dumps(comparison, ensure_ascii=False, indent=2)
        
    except Exception as e:
        return json.dumps({
            "error": f"比较文档时出错: {str(e)}"
        }, ensure_ascii=False)


# 辅助函数
def _extract_core_properties(doc):
    """提取文档的核心属性"""
    props = {}
    core_props = doc.core_properties
    
    if core_props.title:
        props["title"] = core_props.title
    if core_props.author:
        props["author"] = core_props.author
    if core_props.subject:
        props["subject"] = core_props.subject
    if core_props.keywords:
        props["keywords"] = core_props.keywords
    if core_props.category:
        props["category"] = core_props.category
    if core_props.created:
        props["created"] = str(core_props.created)
    if core_props.modified:
        props["modified"] = str(core_props.modified)
    
    return props


def _extract_styles(doc):
    """提取文档的样式信息"""
    styles = []
    
    for style in doc.styles:
        style_info = {
            "name": style.name,
            "type": str(style.type) if hasattr(style, 'type') else "Unknown"
        }
        
        # 提取字体信息
        if hasattr(style, 'font') and style.font:
            font_info = {}
            if style.font.name:
                font_info["name"] = style.font.name
            if style.font.size:
                font_info["size"] = str(style.font.size)
            if style.font.bold is not None:
                font_info["bold"] = style.font.bold
            if style.font.italic is not None:
                font_info["italic"] = style.font.italic
            if style.font.underline is not None:
                font_info["underline"] = style.font.underline
            if style.font.color and style.font.color.rgb:
                font_info["color"] = str(style.font.color.rgb)
            
            if font_info:
                style_info["font"] = font_info
        
        styles.append(style_info)
    
    return styles


def _extract_images(doc):
    """提取文档中的图片"""
    images = []
    image_parts = []
    
    # 收集文档中的图片部分
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_parts.append((rel.target_ref, rel._target))
    
    # 提取图片信息
    for i, (image_id, image_part) in enumerate(image_parts):
        try:
            # 获取图片数据
            image_data = image_part.blob
            image_type = image_part.content_type.split('/')[-1]  # 如 'image/jpeg' -> 'jpeg'
            
            # Base64编码图片数据
            base64_data = base64.b64encode(image_data).decode('utf-8')
            
            images.append({
                "id": f"img{i}",
                "image_id": image_id,
                "content_type": image_part.content_type,
                "size": len(image_data),
                "base64_data": f"data:{image_part.content_type};base64,{base64_data}"
            })
        except:
            # 忽略无法处理的图片
            pass
    
    return images


def _parse_measurement(value):
    """解析测量值（如页边距、页面大小等）"""
    if isinstance(value, (int, float)):
        return value
    
    if isinstance(value, str):
        # 移除单位
        value = value.lower()
        if "pt" in value:
            return Pt(float(value.replace("pt", "").strip()))
        elif "inch" in value:
            return Inches(float(value.replace("inch", "").strip()))
        else:
            # 尝试直接转换为浮点数
            try:
                return float(value)
            except:
                pass
    
    # 默认返回1英寸
    return Inches(1)


def _apply_styles(doc, styles_data):
    """应用样式到文档"""
    for style_data in styles_data:
        style_name = style_data.get("name", "")
        if not style_name:
            continue
        
        # 检查样式是否存在
        if style_name in doc.styles:
            style = doc.styles[style_name]
        else:
            # 创建新样式
            try:
                style_type_map = {
                    "PARAGRAPH": WD_STYLE_TYPE.PARAGRAPH,
                    "CHARACTER": WD_STYLE_TYPE.CHARACTER,
                    "TABLE": WD_STYLE_TYPE.TABLE,
                    "LIST": WD_STYLE_TYPE.LIST
                }
                
                style_type = style_data.get("type", "PARAGRAPH")
                if style_type in style_type_map:
                    style = doc.styles.add_style(style_name, style_type_map[style_type])
                else:
                    style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            except:
                continue  # 如果无法创建样式，跳过
        
        # 应用字体设置
        if "font" in style_data and hasattr(style, 'font'):
            font_data = style_data["font"]
            if "name" in font_data:
                style.font.name = font_data["name"]
            if "size" in font_data:
                try:
                    size_value = font_data["size"]
                    if isinstance(size_value, str) and "Pt" in size_value:
                        size_pt = float(size_value.replace("Pt", "").strip())
                        style.font.size = Pt(size_pt)
                    else:
                        style.font.size = Pt(float(size_value))
                except:
                    pass  # 忽略无效的字体大小
            if "bold" in font_data:
                style.font.bold = font_data["bold"]
            if "italic" in font_data:
                style.font.italic = font_data["italic"]
            if "underline" in font_data:
                style.font.underline = font_data["underline"]
            if "color" in font_data:
                try:
                    color_str = font_data["color"]
                    if color_str.startswith("RGB"):
                        rgb_values = color_str.replace("RGB", "").replace("(", "").replace(")", "").split(",")
                        if len(rgb_values) == 3:
                            r, g, b = map(int, rgb_values)
                            style.font.color.rgb = RGBColor(r, g, b)
                except:
                    pass  # 忽略无效的颜色